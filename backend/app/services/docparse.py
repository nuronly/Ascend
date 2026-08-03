"""文档解析（PLAN §3.5）。

**难点不在翻译，在切段。**

PDF 抽出来的文本天然是碎的：一行一个片段、跨页断句、双栏交错、
连字符断词。直接按行翻译会得到一堆语义不完整的碎片。
所以这里的重点是「把碎片重新拼回段落」的启发式合并。

许可证提醒（PLAN §3.5）：
  ❌ 不用 PyMuPDF —— AGPL-3.0，对外提供网络服务就要开源全部代码
  ✅ pdfplumber (MIT) / pypdfium2 (Apache) / python-docx (MIT) / ebooklib
"""

from __future__ import annotations

import hashlib
import io
import logging
import re
from dataclasses import dataclass, field

log = logging.getLogger(__name__)

MAX_BLOCK = 1800  # 单段上限，太长的段落切开以免翻译超时


@dataclass
class Block:
    text: str
    page: int = 0
    block_type: str = "paragraph"
    bbox: dict = field(default_factory=dict)

    @property
    def hash(self) -> str:
        return hashlib.sha256(self.text.strip().encode()).hexdigest()


# ─────────────────────────────────────────────────────────────
# 文本清洗与段落合并
# ─────────────────────────────────────────────────────────────
_HEADING = re.compile(
    r"^(?:\d+(?:\.\d+)*\.?\s+\S|第[一二三四五六七八九十百]+[章节部分]|"
    r"(?:Chapter|Section|Part|Appendix)\s+[\dIVXA-Z]+|摘要|Abstract|"
    r"References|Bibliography|参考文献|致谢|Acknowledg(?:e)?ments?)",
    re.I,
)
_LIST_ITEM = re.compile(r"^\s*(?:[-•·▪–—*]|\(?\d{1,2}[.)]|\(?[a-z][.)])\s+")
_PAGE_NUM = re.compile(r"^\s*(?:[-–—]\s*)?\d{1,4}\s*(?:[-–—])?\s*$")
_URL_ONLY = re.compile(r"^\s*(?:https?://|www\.)\S+\s*$", re.I)


def _looks_like_heading(line: str) -> bool:
    s = line.strip()
    if not s or len(s) > 120:
        return False
    return bool(_HEADING.match(s))


def _ends_sentence(s: str) -> bool:
    return bool(re.search(r"[.!?。！？；;:：)\]}”』】]\s*$", s.rstrip()))


def _join(prev: str, nxt: str) -> str:
    """跨行拼接。

    三个坑：
      · 英文断词连字符 `informa-\ntion` 要还原成 `information`
      · 中文之间不能插空格
      · 英文之间必须有空格
    """
    prev = prev.rstrip()
    nxt = nxt.lstrip()
    if prev.endswith("-") and re.match(r"^[a-z]", nxt):
        return prev[:-1] + nxt
    if not prev or not nxt:
        return prev + nxt
    if re.search(r"[\u4e00-\u9fff]$", prev) or re.match(r"^[\u4e00-\u9fff]", nxt):
        return prev + nxt
    return prev + " " + nxt


def merge_lines(lines: list[tuple[str, int]]) -> list[Block]:
    """把「行」合并成「段」。这是整个文档模式质量的分水岭。

    判定规则（按优先级）：
      1. 空行 → 段落结束
      2. 标题行 → 自成一段
      3. 列表项开头 → 新段
      4. 上一行以句末标点结尾 且 本行像新句子开头 → 新段
      5. 否则 → 续接上一段（这是 PDF 里最常见的情况）
    """
    blocks: list[Block] = []
    buf = ""
    buf_page = 0
    buf_type = "paragraph"

    def flush():
        nonlocal buf, buf_type
        t = buf.strip()
        if t and not _PAGE_NUM.match(t):
            for piece in _split_long(t):
                blocks.append(Block(text=piece, page=buf_page, block_type=buf_type))
        buf = ""
        buf_type = "paragraph"

    for raw, page in lines:
        line = raw.rstrip()
        stripped = line.strip()

        if not stripped:
            flush()
            continue
        if _PAGE_NUM.match(stripped) or _URL_ONLY.match(stripped):
            # 页码、孤立 URL 属于版式噪音，不值得翻译
            flush()
            continue

        if _looks_like_heading(stripped):
            flush()
            buf, buf_page, buf_type = stripped, page, "heading"
            flush()
            continue

        if not buf:
            buf, buf_page = stripped, page
            buf_type = "list" if _LIST_ITEM.match(line) else "paragraph"
            continue

        # 新的列表项 → 另起一段
        if _LIST_ITEM.match(line):
            flush()
            buf, buf_page, buf_type = stripped, page, "list"
            continue

        # 上一段已经结句，且本行是新句开头（大写字母 / 中文 / 数字编号）
        if _ends_sentence(buf) and re.match(r"^[A-Z\u4e00-\u9fff\d]", stripped):
            # 但如果上一段很短，多半是被硬切开的标题或残句，仍然接上
            if len(buf) > 60:
                flush()
                buf, buf_page = stripped, page
                continue

        buf = _join(buf, stripped)

    flush()
    return blocks


def _split_long(text: str) -> list[str]:
    """超长段落按句子边界切开，避免单次翻译请求过大。"""
    if len(text) <= MAX_BLOCK:
        return [text]
    parts: list[str] = []
    cur = ""
    for sent in re.split(r"(?<=[.!?。！？])\s*", text):
        if not sent:
            continue
        if len(cur) + len(sent) > MAX_BLOCK and cur:
            parts.append(cur)
            cur = sent
        else:
            cur += sent
    if cur:
        parts.append(cur)
    return parts


# ─────────────────────────────────────────────────────────────
# PDF
# ─────────────────────────────────────────────────────────────
def parse_pdf(data: bytes) -> tuple[list[Block], int]:
    """用 pdfplumber 抽文本层。

    双栏论文是已知的软肋（PLAN §7 风险 #5）：这里按 x 坐标做了
    简单的分栏检测，能救回大部分两栏排版；复杂版式仍会有错位，
    所以 arXiv 一律优先走 HTML 版。
    """
    import pdfplumber

    lines: list[tuple[str, int]] = []
    page_count = 0

    with pdfplumber.open(io.BytesIO(data)) as pdf:
        page_count = len(pdf.pages)
        for pno, page in enumerate(pdf.pages, start=1):
            words = page.extract_words(use_text_flow=False, keep_blank_chars=False)
            if not words:
                continue

            width = page.width or 612
            # 分栏检测：统计词的中心 x，若中间地带明显空缺则判为双栏
            mid = width / 2
            left = sum(1 for w in words if (w["x0"] + w["x1"]) / 2 < mid * 0.92)
            right = sum(1 for w in words if (w["x0"] + w["x1"]) / 2 > mid * 1.08)
            middle = len(words) - left - right
            two_col = left > 25 and right > 25 and middle < len(words) * 0.12

            groups = (
                [
                    [w for w in words if (w["x0"] + w["x1"]) / 2 < mid],
                    [w for w in words if (w["x0"] + w["x1"]) / 2 >= mid],
                ]
                if two_col
                else [words]
            )

            for col in groups:
                if not col:
                    continue
                # 按行聚类：top 坐标接近的算同一行
                col = sorted(col, key=lambda w: (round(w["top"], 1), w["x0"]))
                cur_top = None
                cur: list[str] = []
                for w in col:
                    if cur_top is None or abs(w["top"] - cur_top) <= 2.6:
                        cur.append(w["text"])
                        cur_top = w["top"] if cur_top is None else cur_top
                    else:
                        lines.append((" ".join(cur), pno))
                        cur = [w["text"]]
                        cur_top = w["top"]
                if cur:
                    lines.append((" ".join(cur), pno))
                lines.append(("", pno))  # 栏与栏之间断开

    return merge_lines(lines), page_count


# ─────────────────────────────────────────────────────────────
# docx / epub / markdown / txt
# ─────────────────────────────────────────────────────────────
def parse_docx(data: bytes) -> tuple[list[Block], int]:
    import docx

    doc = docx.Document(io.BytesIO(data))
    blocks: list[Block] = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        style = (p.style.name or "").lower()
        btype = "heading" if "heading" in style or "title" in style else "paragraph"
        for piece in _split_long(t):
            blocks.append(Block(text=piece, page=1, block_type=btype))

    # 表格按行展开，保留内容但不强求版式
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                blocks.append(Block(text=" | ".join(cells), page=1, block_type="table"))
    return blocks, 1


def parse_epub(data: bytes) -> tuple[list[Block], int]:
    import ebooklib
    from bs4 import BeautifulSoup
    from ebooklib import epub

    import tempfile, os

    # ebooklib 只接受文件路径
    fd, path = tempfile.mkstemp(suffix=".epub")
    try:
        with os.fdopen(fd, "wb") as f:
            f.write(data)
        book = epub.read_epub(path)
        blocks: list[Block] = []
        page = 0
        for item in book.get_items():
            if item.get_type() != ebooklib.ITEM_DOCUMENT:
                continue
            page += 1
            soup = BeautifulSoup(item.get_content(), "lxml")
            blocks.extend(_from_soup(soup, page))
        return blocks, max(page, 1)
    finally:
        try:
            os.unlink(path)
        except OSError:
            pass


def parse_markdown(text: str) -> tuple[list[Block], int]:
    blocks: list[Block] = []
    in_code = False
    buf: list[str] = []
    btype = "paragraph"

    def flush():
        nonlocal buf, btype
        t = "\n".join(buf).strip()
        if t:
            for piece in _split_long(t) if btype != "code" else [t]:
                blocks.append(Block(text=piece, page=1, block_type=btype))
        buf = []
        btype = "paragraph"

    for line in text.splitlines():
        if line.lstrip().startswith("```"):
            if in_code:
                buf.append(line)
                flush()
                in_code = False
            else:
                flush()
                in_code = True
                btype = "code"
                buf.append(line)
            continue
        if in_code:
            buf.append(line)
            continue
        if not line.strip():
            flush()
            continue
        if line.lstrip().startswith("#"):
            flush()
            buf = [line.strip()]
            btype = "heading"
            flush()
            continue
        buf.append(line.strip())
    flush()
    return blocks, 1


def parse_text(text: str) -> tuple[list[Block], int]:
    return merge_lines([(ln, 1) for ln in text.splitlines()]), 1


def _from_soup(soup, page: int) -> list[Block]:
    from bs4 import BeautifulSoup  # noqa: F401

    blocks: list[Block] = []
    # 去掉不该翻译的噪音
    for tag in soup(["script", "style", "nav", "footer", "noscript"]):
        tag.decompose()

    for el in soup.find_all(
        ["h1", "h2", "h3", "h4", "h5", "h6", "p", "li", "pre", "blockquote", "figcaption"]
    ):
        # 跳过嵌套元素，避免同一段被抓两遍
        if el.find_parent(["li", "pre", "blockquote"]) and el.name not in ("pre",):
            continue
        t = el.get_text(" ", strip=True)
        if not t or len(t) < 2:
            continue
        name = el.name
        btype = (
            "heading"
            if name.startswith("h")
            else "code"
            if name == "pre"
            else "list"
            if name == "li"
            else "paragraph"
        )
        for piece in _split_long(t) if btype != "code" else [t]:
            blocks.append(Block(text=piece, page=page, block_type=btype))
    return blocks


def parse_html(html: str) -> tuple[list[Block], int]:
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(html, "lxml")
    # arXiv 的 HTML 版正文通常在 <article> 或 .ltx_page_content 里
    root = (
        soup.find("article")
        or soup.find(class_="ltx_page_content")
        or soup.find("main")
        or soup
    )
    return _from_soup(root, 1), 1


def title_from_html(html: str) -> str:
    from bs4 import BeautifulSoup

    soup = BeautifulSoup(html, "lxml")
    for sel in ("h1.ltx_title", "h1", "title"):
        el = soup.select_one(sel)
        if el:
            t = el.get_text(" ", strip=True)
            if t:
                return t[:400]
    return ""


# ─────────────────────────────────────────────────────────────
# 统一入口
# ─────────────────────────────────────────────────────────────
def parse(data: bytes, filename: str, mime: str = "") -> tuple[list[Block], int]:
    name = filename.lower()
    if name.endswith(".pdf") or "pdf" in mime:
        return parse_pdf(data)
    if name.endswith(".docx"):
        return parse_docx(data)
    if name.endswith(".epub"):
        return parse_epub(data)

    text = data.decode("utf-8", errors="replace")
    if name.endswith((".md", ".markdown")):
        return parse_markdown(text)
    if name.endswith((".html", ".htm")):
        return parse_html(text)
    return parse_text(text)


SUPPORTED_EXT = (".pdf", ".docx", ".epub", ".md", ".markdown", ".txt", ".html", ".htm")
